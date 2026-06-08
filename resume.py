"""
job_pipeline/resume.py
======================
Generates a job-tailored version of the base resume using the Claude API.
Only invoked for strong matches (see run_pipeline.py) — Tier A/B roles scoring
>= 70%. The base resume lives in the BASE_RESUME/ folder (.docx preferred).

Output is saved as Markdown under applications/<Company>/, mirroring how cover
letters are stored.
"""

import os
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import anthropic
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

BASE_DIR = Path(__file__).parent
BASE_RESUME_DIR = BASE_DIR / "BASE_RESUME"

# The user's resume-editing prompt, used verbatim as the system instruction.
RESUME_SYSTEM_PROMPT = """You are an expert resume editor specializing in Data Engineering / Data Science roles.

I will provide:
1. My base resume
2. A job description

Your task is to tailor my resume specifically for this job while following STRICT constraints:

-----------------------
HARD CONSTRAINTS
-----------------------
- Keep the overall structure EXACTLY the same (Summary, Skills, Experience, Projects, Education).
- Do NOT add new sections.
- Do NOT remove existing roles or projects.
- Resume must remain within ONE PAGE.
- Avoid adding fluff or generic statements.
- No fake experience — only reframe and optimize existing content.
- Maintain realistic, explainable claims aligned with my background.

-----------------------
OBJECTIVE
-----------------------
Optimize my resume to:
- Maximize alignment with the job description
- Improve ATS keyword matching
- Highlight most relevant skills and experience
- Make it easy to justify everything in interviews

-----------------------
WHAT TO MODIFY
-----------------------

1. SUMMARY SECTION
- Rewrite to align strongly with the job role
- Include key tools/skills from JD (without keyword stuffing)
- Keep it concise (3–4 lines max)
- Focus on impact + specialization

2. SKILLS SECTION
- Reorder skills based on relevance to JD
- Add missing skills ONLY if they can be justified from experience/projects
- Group skills logically (e.g., Programming, ML, Cloud, BI)
- Remove less relevant skills if needed

3. EXPERIENCE SECTION
- Rephrase bullet points to reflect:
  - Impact (metrics-driven)
  - Relevance to JD (tools, techniques, domain)
- Prioritize most relevant bullets at top
- Adjust wording to match JD terminology
- Highlight:
  - Data pipelines / ML / analytics / BI depending on role
- Keep 4–6 bullets per role max

4. PROJECTS SECTION
- Emphasize projects most relevant to job
- Reorder projects if needed
- Adjust descriptions to highlight:
  - tools from JD
  - business impact
  - ML/analytics/system design relevance

5. KEYWORD ALIGNMENT
- Ensure important keywords from JD are naturally included:
  (e.g., ETL, ML models, Power BI, SQL, pipelines, etc.)
- Avoid unnatural keyword stuffing

6. SPACE OPTIMIZATION
- Ensure no excessive white space
- Keep content dense but readable
- Use concise bullet points

-----------------------
OUTPUT FORMAT
-----------------------
- Return FULL updated resume
- Keep formatting clean and consistent
- Do NOT explain changes
- Do NOT add comments
- Only output final resume"""


def _docx_to_text(path: Path) -> str:
    """Extract plain text from a .docx. Uses python-docx if available, else
    falls back to parsing word/document.xml directly (stdlib only)."""
    try:
        import docx  # python-docx, if installed
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(p for p in parts if p.strip())
    except ImportError:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        with zipfile.ZipFile(path) as z:
            root = ET.fromstring(z.read("word/document.xml"))
        lines = []
        for para in root.iter(ns + "p"):
            line = "".join(t.text for t in para.iter(ns + "t") if t.text).strip()
            if line:
                lines.append(line)
        return "\n".join(lines)


def load_base_resume() -> str:
    """Load the base resume text from the BASE_RESUME/ folder.

    Prefers .docx, then .md/.txt. Returns "" if no resume file is found.
    """
    if not BASE_RESUME_DIR.exists():
        return ""
    for pattern in ("*.docx", "*.md", "*.txt"):
        for path in sorted(BASE_RESUME_DIR.glob(pattern)):
            if path.name.startswith("~$"):  # skip Word lock files
                continue
            if path.suffix == ".docx":
                text = _docx_to_text(path)
            else:
                text = path.read_text(encoding="utf-8")
            if text.strip():
                return text
    return ""


def generate_tailored_resume(base_resume: str, company_name: str, job_title: str,
                             job_description: str) -> str:
    """Return a job-tailored version of the base resume (Markdown text)."""
    user_message = f"""JOB DESCRIPTION:
{job_description[:4000]}

BASE RESUME:
{base_resume[:6000]}"""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=3000,
        system=RESUME_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )

    resume_body = response.content[0].text.strip()
    today = datetime.now().strftime("%B %d, %Y")

    # Light metadata header for file traceability (mirrors cover letters).
    return (f"<!-- Tailored resume | {company_name} — {job_title} | "
            f"Generated {today} -->\n\n{resume_body}\n")
